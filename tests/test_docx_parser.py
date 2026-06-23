import io
import unittest

from docx import Document as DocxDocument

from utils.docx_parser import parse_docx_content


def _make_docx(paragraphs: list[str], table_rows: list[list[str]] | None = None) -> bytes:
    doc = DocxDocument()
    for text in paragraphs:
        doc.add_paragraph(text)
    if table_rows:
        table = doc.add_table(rows=len(table_rows), cols=len(table_rows[0]))
        for i, row_data in enumerate(table_rows):
            for j, cell_text in enumerate(row_data):
                table.cell(i, j).text = cell_text
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


class TestParseDocxContent(unittest.TestCase):
    def test_extracts_paragraphs(self) -> None:
        docx_bytes = _make_docx(["Introduction", "Main body text", "Conclusion"])
        result = parse_docx_content(docx_bytes)
        self.assertIn("Introduction", result)
        self.assertIn("Main body text", result)
        self.assertIn("Conclusion", result)

    def test_extracts_table_cells(self) -> None:
        docx_bytes = _make_docx([], table_rows=[["Header A", "Header B"], ["Cell 1", "Cell 2"]])
        result = parse_docx_content(docx_bytes)
        self.assertIn("Header A", result)
        self.assertIn("Cell 2", result)

    def test_skips_empty_paragraphs(self) -> None:
        docx_bytes = _make_docx(["", "Non-empty", ""])
        result = parse_docx_content(docx_bytes)
        self.assertIn("Non-empty", result)
        self.assertNotIn("\n\n\n", result)

    def test_handles_invalid_bytes(self) -> None:
        result = parse_docx_content(b"not a docx file")
        self.assertTrue(result.startswith("Error extracting text from DOCX:"))


if __name__ == "__main__":
    unittest.main()
